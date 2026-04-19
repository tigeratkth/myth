"""文件读写工具

职责：
  1. 从上传文件 / 路径统一抽取文本（txt / md / docx / pdf）
  2. 把一个完整 WorkflowState 落盘为 outputs/{task_id}/ 下的 JSON + Markdown
  3. 列举、读取、删除历史任务
"""
from __future__ import annotations

import json
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, BinaryIO, Iterable

from .state import MODULE_KEYS, MODULE_LABELS, WorkflowState


# ============================================================
# 一、原始资料读取
# ============================================================
def read_material(file: BinaryIO | str | Path, filename: str | None = None) -> str:
    """从 txt / md / docx / pdf 提取纯文本。

    - file 可以是路径字符串、Path 或二进制流（如 Streamlit 的 UploadedFile）。
    - filename 用于判断后缀；若 file 是路径则可省略。
    """
    if isinstance(file, (str, Path)):
        path = Path(file)
        name = filename or path.name
        ext = path.suffix.lower().lstrip(".")
        data = path.read_bytes()
    else:
        if not filename:
            raise ValueError("传入二进制流时必须同时提供 filename")
        name = filename
        ext = Path(filename).suffix.lower().lstrip(".")
        data = file.read()
        # 重置指针，便于调用方后续再次使用
        try:
            file.seek(0)
        except Exception:  # noqa: BLE001
            pass

    if ext in ("txt", "md", ""):
        return _decode_text(data)
    if ext == "docx":
        return _read_docx(data)
    if ext == "pdf":
        return _read_pdf(data)
    raise ValueError(f"不支持的文件类型：{name}（仅支持 txt/md/docx/pdf）")


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_docx(data: bytes) -> str:
    import io
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(data: bytes) -> str:
    import io
    from pypdf import PdfReader  # type: ignore

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    return "\n\n".join(pages).strip()


# ============================================================
# 二、任务落盘 / 读取
# ============================================================
def save_task(state: WorkflowState, outputs_dir: str | Path = "outputs") -> str:
    """把完整 state 保存到 outputs/{task_id}/ 下。

    产物：
      - meta.json   : 元信息（历史列表扫描它）
      - state.json  : 完整 state（用于"继续"或"基于此任务重跑"）
      - m1.md..m4.md: 每模块的 Markdown 渲染
      - report.md   : 合并的完整报告

    返回任务目录的绝对路径字符串。
    """
    task_id = state.get("task_id") or _fallback_task_id()
    out_root = Path(outputs_dir)
    out_root.mkdir(parents=True, exist_ok=True)
    task_dir = out_root / task_id
    task_dir.mkdir(parents=True, exist_ok=True)

    (task_dir / "meta.json").write_text(
        json.dumps(_build_meta(state), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (task_dir / "state.json").write_text(
        json.dumps(_serializable_state(state), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    module_sections: list[str] = []
    for key in MODULE_KEYS:
        md = render_module_markdown(state, key)
        (task_dir / f"{key}.md").write_text(md, encoding="utf-8")
        module_sections.append(md)

    report_md = _render_report(state, module_sections)
    (task_dir / "report.md").write_text(report_md, encoding="utf-8")

    return str(task_dir.resolve())


def load_task(task_id: str, outputs_dir: str | Path = "outputs") -> WorkflowState:
    """反序列化 outputs/{task_id}/state.json 为 WorkflowState。"""
    path = Path(outputs_dir) / task_id / "state.json"
    if not path.exists():
        raise FileNotFoundError(f"任务不存在：{path}")
    data = json.loads(path.read_text(encoding="utf-8"))
    return WorkflowState(**data)  # TypedDict 构造


def list_tasks(outputs_dir: str | Path = "outputs") -> list[dict]:
    """扫描所有 meta.json，按 created_at 倒序返回。"""
    root = Path(outputs_dir)
    if not root.exists():
        return []
    rows: list[dict] = []
    for sub in root.iterdir():
        if not sub.is_dir():
            continue
        meta_path = sub / "meta.json"
        if not meta_path.exists():
            continue
        try:
            rows.append(json.loads(meta_path.read_text(encoding="utf-8")))
        except Exception:  # noqa: BLE001
            continue
    rows.sort(key=lambda r: r.get("created_at", ""), reverse=True)
    return rows


def delete_task(task_id: str, outputs_dir: str | Path = "outputs") -> None:
    """递归删除 outputs/{task_id}/。"""
    task_dir = Path(outputs_dir) / task_id
    if task_dir.exists():
        shutil.rmtree(task_dir)


# ============================================================
# 三、辅助：meta / markdown 渲染
# ============================================================
def _fallback_task_id() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S-unnamed")


def _serializable_state(state: WorkflowState) -> dict:
    """TypedDict → dict，并剔除不可序列化的字段（目前没有，保留口子）。"""
    return dict(deepcopy(state))


def _build_meta(state: WorkflowState) -> dict:
    """构造 outputs/{task_id}/meta.json 的内容（前端历史列表直接用）。"""
    models = state.get("module_models") or {}
    metas = {k: state.get(f"{k}_meta", {}) or {} for k in MODULE_KEYS}
    # 计算总耗时 = 各模块 duration_ms 之和（有就加）
    total_duration = sum(int(m.get("duration_ms") or 0) for m in metas.values())

    completed_at: str | None = None
    status = state.get("overall_status") or "idle"
    if status == "completed":
        completed_at = metas.get("m4", {}).get("finished_at")

    return {
        "task_id": state.get("task_id", ""),
        "task_name": state.get("task_name", ""),
        "mode": state.get("mode", "auto"),
        "status": status,
        "created_at": state.get("created_at", ""),
        "completed_at": completed_at,
        "duration_ms": total_duration,
        "models": {k: models.get(k, "") for k in MODULE_KEYS},
        "current_stage": state.get("current_module", "m1"),
        "source_type": state.get("source_type", "paste"),
        "source_filename": state.get("source_filename", ""),
        "tokens": {
            k: {
                "in": int(metas[k].get("tokens_in") or 0),
                "out": int(metas[k].get("tokens_out") or 0),
            }
            for k in MODULE_KEYS
        },
    }


# ---------- Markdown 渲染 ----------
def render_module_markdown(state: WorkflowState, module_key: str) -> str:
    """把单个模块的输出渲染成 Markdown（不是原始 JSON）。"""
    label = MODULE_LABELS.get(module_key, module_key)  # type: ignore[arg-type]
    output = state.get(f"{module_key}_output") or {}
    meta = state.get(f"{module_key}_meta") or {}

    lines: list[str] = [f"# 模块 {module_key.upper()} · {label}", ""]
    if meta:
        status = meta.get("status", "unknown")
        model = meta.get("model", "-")
        duration = meta.get("duration_ms")
        dur_str = f"{duration}ms" if duration else "-"
        lines.append(f"> 状态：`{status}` · 模型：`{model}` · 用时：`{dur_str}`")
        if meta.get("error"):
            lines.append(f"> 错误：{meta['error']}")
        lines.append("")

    if not output:
        lines.append("_（暂无输出）_")
        return "\n".join(lines)

    renderer = {
        "m1": _render_m1,
        "m2": _render_m2,
        "m3": _render_m3,
        "m4": _render_m4,
    }.get(module_key)
    if renderer:
        lines.extend(renderer(output))
    else:
        lines.append("```json")
        lines.append(json.dumps(output, ensure_ascii=False, indent=2))
        lines.append("```")
    return "\n".join(lines)


def _render_m1(out: dict) -> list[str]:
    lines: list[str] = ["## 技术要点", ""]
    for i, tp in enumerate(out.get("tech_points", []) or [], start=1):
        lines.append(f"### {i}. {tp.get('plain', '-')}")
        if tp.get("original"):
            lines.append(f"- **原文表述**：{tp['original']}")
        if tp.get("analogy"):
            lines.append(f"- **类比**：{tp['analogy']}")
        if tp.get("params"):
            lines.append(f"- **参数**：{tp['params']}")
        lines.append("")
    lines += ["## 关键能力", ""]
    lines += [f"- {c}" for c in out.get("capabilities", []) or []] or ["_无_"]
    lines += ["", "## 能力边界 / 前提", ""]
    lines += [f"- {b}" for b in out.get("boundaries", []) or []] or ["_无_"]
    return lines


def _render_m2(out: dict) -> list[str]:
    lines: list[str] = []
    if out.get("core_claim"):
        lines += ["## 核心技术IP 主张", "", f"> {out['core_claim']}", ""]
    lines += ["## 技术卖点", ""]
    for i, sp in enumerate(out.get("selling_points", []) or [], start=1):
        lines.append(f"### {i}. {sp.get('name', '-')}")
        if sp.get("user_description"):
            lines.append(f"- **用户语言描述**：{sp['user_description']}")
        if sp.get("user_need"):
            lines.append(f"- **对应用户需求**：{sp['user_need']}")
        if sp.get("supporting_points"):
            lines.append(f"- **支撑技术要点**：{', '.join(sp['supporting_points'])}")
        if sp.get("differentiation"):
            lines.append(f"- **差异化**：{sp['differentiation']}")
        lines.append("")
    lines += ["## 目标用户假设", ""]
    lines += [f"- {u}" for u in out.get("target_user_hypothesis", []) or []] or ["_无_"]
    lines += ["", "## 不建议对外强调的技术点", ""]
    lines += [f"- {f}" for f in out.get("filtered_points", []) or []] or ["_无_"]
    return lines


def _render_m3(out: dict) -> list[str]:
    lines: list[str] = ["## 目标人群画像", ""]
    for i, au in enumerate(out.get("target_audiences", []) or [], start=1):
        lines.append(f"### {i}. {au.get('segment', '-')}")
        if au.get("description"):
            lines.append(f"- **画像**：{au['description']}")
        if au.get("pain_points"):
            lines.append("- **痛点**：")
            lines += [f"  - {p}" for p in au["pain_points"]]
        if au.get("preferred_channels"):
            lines.append(f"- **常出现渠道**：{', '.join(au['preferred_channels'])}")
        lines.append("")
    lines += ["## 核心渠道", ""]
    lines += [f"- {c}" for c in out.get("core_channels", []) or []] or ["_无_"]
    lines += ["", "## 内容矩阵", ""]
    phase_label = {"awareness": "认知", "seeding": "种草", "conversion": "转化"}
    for cm in out.get("content_matrix", []) or []:
        phase = cm.get("phase", "")
        lines.append(f"### {phase_label.get(phase, phase)}（{phase}）")
        if cm.get("key_selling_points"):
            lines.append(f"- **主打卖点**：{', '.join(cm['key_selling_points'])}")
        if cm.get("content_types"):
            lines.append(f"- **内容形态**：{', '.join(cm['content_types'])}")
        if cm.get("sample_angles"):
            lines.append("- **示例切入角度**：")
            lines += [f"  - {a}" for a in cm["sample_angles"]]
        lines.append("")
    lines += ["## 推广节奏", ""]
    lines += [f"- {p}" for p in out.get("phases", []) or []] or ["_无_"]
    lines += ["", "## 关键 KPI", ""]
    lines += [f"- {k}" for k in out.get("kpis", []) or []] or ["_无_"]
    return lines


def _render_m4(out: dict) -> list[str]:
    lines: list[str] = []
    if out.get("campaign_theme"):
        lines += [f"## 活动主题", "", out["campaign_theme"], ""]
    if out.get("slogan"):
        lines += [f"## Slogan", "", f"> {out['slogan']}", ""]

    vids = out.get("video_scripts", []) or []
    if vids:
        lines += ["## 短视频脚本", ""]
        for i, v in enumerate(vids, start=1):
            lines.append(f"### {i}. {v.get('title', '-')} （{v.get('duration_sec', '?')}s）")
            if v.get("hook"):
                lines.append(f"- **开场钩子**：{v['hook']}")
            if v.get("body"):
                lines.append(f"- **主体**：{v['body']}")
            if v.get("cta"):
                lines.append(f"- **行动号召**：{v['cta']}")
            lines.append("")

    if out.get("article"):
        lines += ["## 公众号长文", "", out["article"], ""]

    posts = out.get("social_posts", []) or []
    if posts:
        lines += ["## 社媒短图文", ""]
        for i, p in enumerate(posts, start=1):
            lines.append(f"### {i}. [{p.get('platform', '-')}] {p.get('title', '-')}")
            if p.get("body"):
                lines.append(p["body"])
            if p.get("hashtags"):
                lines.append("")
                lines.append(" ".join(p["hashtags"]))
            lines.append("")

    posters = out.get("posters", []) or []
    if posters:
        lines += ["## 海报文案", ""]
        for i, pst in enumerate(posters, start=1):
            lines.append(f"### {i}. {pst.get('headline', '-')}")
            if pst.get("subline"):
                lines.append(f"- **副标**：{pst['subline']}")
            if pst.get("visual_keywords"):
                lines.append(f"- **视觉关键词**：{', '.join(pst['visual_keywords'])}")
            lines.append("")

    ev = out.get("offline_event") or {}
    if ev:
        lines += ["## 线下活动方案", ""]
        if ev.get("theme"):
            lines.append(f"- **主题**：{ev['theme']}")
        if ev.get("flow"):
            lines.append("- **流程**：")
            lines += [f"  - {f}" for f in ev["flow"]]
        if ev.get("materials"):
            lines.append(f"- **物料**：{', '.join(ev['materials'])}")
        if ev.get("budget_framework"):
            lines.append(f"- **预算框架**：{ev['budget_framework']}")
        lines.append("")
    return lines


def _render_report(state: WorkflowState, module_sections: Iterable[str]) -> str:
    """把四个模块合并为一份完整报告。"""
    header = [
        f"# {state.get('task_name') or state.get('task_id') or '业务智能体报告'}",
        "",
        f"- 任务 ID：`{state.get('task_id', '-')}`",
        f"- 创建时间：{state.get('created_at', '-')}",
        f"- 运行模式：{state.get('mode', '-')}",
        f"- 整体状态：{state.get('overall_status', '-')}",
        "",
        "---",
        "",
    ]
    body = "\n\n---\n\n".join(module_sections)
    return "\n".join(header) + body + "\n"


__all__ = [
    "read_material",
    "save_task",
    "load_task",
    "list_tasks",
    "delete_task",
    "render_module_markdown",
]
